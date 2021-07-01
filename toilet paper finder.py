import time
import signal
import docx
from docx import Document
import re
import PyDictionary
import json
from bs4 import BeautifulSoup
import requests
import datetime
from selenium import webdriver
from selenium.webdriver.support.ui import Select
from selenium import common
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_UNDERLINE
# from selenium import JavascriptExecutor




# TODO: Pull Unique list of terms to look for
final_list_of_sentences = []
test = True
text_file = open("keywords.txt", "r")
keywords = text_file.read()
# keywords = json.loads(keywords)
# skippable = keywords['skippable']
# possitive_verb = keywords['possitive_verb']
# negative_verb = keywords['negative_verb']
# possitive_noun = keywords['possitive_noun']
# negative_noun = keywords['negative_noun']
# possitive_adjective = keywords['possitive_adjective']
# negative_adjective = keywords['negative_adjective']
# possitive_adverb = keywords['possitive_adverb']
# negative_adverb = keywords['negative_adverb']
# conjunctions = keywords['conjunctions']
link_file = open('link_list.txt', 'r')
list_of_links = link_file.read()
list_of_links = json.loads(list_of_links)
company = ''
job_title = ''
skill = ''

# submitted = list_of_links['submitted']



def ask_to_add_word_to_list(word, definition, type):
    print(definition)
    repeat = True
    while repeat:
        answer = input(f'Do you want to add \"{word}\" to the list? ({type})[g/b]')
        if answer in ['G', 'g', 'B', 'b', '', 'skip']:
            repeat = False
        if answer in ['G', 'g']:
            return True
        if answer in ('b', 'B'):
            return False
        if answer in ['skip', '']:
            return 'Skip'

def fill_form(driver):
    successful = False
    failure = 0
    try:
        user_name = driver.find_element_by_xpath('//*[@id="input-applicant.name"]')
        time.sleep(1)
        user_name.send_keys('Shelby Whittaker')
        successful = True
    except Exception:
        failure += 1
    try:
        user_name = driver.find_element_by_xpath('//*[@id="input-applicant.name"]')
        time.sleep(1)
        user_name.send_keys('Shelby Whittaker')
        successful = True
    except Exception:
        failure += 1
    try:
        user_email = driver.find_element_by_xpath('//*[@id="input-applicant.email"]')
        user_email.send_keys('louwhi1992@hotmail.com')
        successful = True
    except Exception:
        failure += 1
    try:
        user_phone_number = driver.find_element_by_xpath('//*[@id="input-applicant.phoneNumber"]')
        user_phone_number.send_keys('8015987752')
        successful = True
    except Exception:
        failure += 1
    try:
        user_resume = driver.find_elements_by_xpath('//*[@id="ia-FilePicker-resume"]')
        user_resume.send_keys(f'C:/Users/Lou/PycharmProjects/job_searcher/Shelby Whittaker {job_title} at {company}.docx')
        user_resume.send_keys(Keys.RETURN)
        successful = True
    except Exception:
        failure += 1
    return successful


def submit_resume(url, driver):
    filename = f"Shelby Whittaker {job_title} at {company}.docx"
    driver.get(url)
    buttons = driver.find_elements_by_css_selector('button')
    # buttons[2].click()
    for span in buttons:
        if span.text == 'Apply Now':
            span.click()
            break
    time.sleep(1)
    iframes = driver.find_elements_by_tag_name('iframe')
    success = False
    while not success:
        for frame in iframes:
            if not success:
                driver.switch_to.default_content()
                driver.switch_to.frame(frame)
                response = fill_form(driver)
                if response:
                    success = True
                    break
                else:
                    iframes2 = driver.find_elements_by_tag_name('iframe')
                    if iframes2:
                        for frame2 in iframes2:
                            driver.switch_to.frame(frame2)
                            response = fill_form(driver)
                            if response:
                                success = True
                                break

    input('submitted')
    #TODO Add information that will manually submit a resume.


def validate_resume(resume_name, url, driver):
    invalid = True
    print(url)
    while invalid:
        if input(f"Please review {resume_name} before continuing. Once you have reviewed the "
                 f"resume and feel it is acceptable, type 'continue'.").lower() == "continue":
            invalid = False
    submit_resume(url, driver)


def determine_resume_content(list_of_sentences, word_list, company, job_title, location, url):
    # for sentence in list_of_sentences:
    #      sentence = sentence.replace(',', '').replace(':', '').replace('\'', '').replace(';', '')
    #      final_list_of_sentences.append(sentence)
    # final_list_of_words = ' '.join(final_list_of_sentences).split(' ')
    # massive_list = possitive_adjective + negative_adjective + possitive_adverb + negative_adverb + possitive_noun + negative_noun + possitive_verb + negative_verb + conjunctions
    # possitive_list = possitive_verb + possitive_noun + possitive_adjective + possitive_adverb
    local_words = []
    # for word in final_list_of_words:
    #     if word.lower() in skippable:
    #         continue
    #     elif word.lower() in massive_list:
    #         if word.lower() in (possitive_adjective or possitive_adverb or possitive_noun or possitive_verb):
    #             if word.lower() not in local_words:
    #                 local_words.append(word.lower())
    #                 continue
    #     # elif PyDictionary.PyDictionary.meaning(word):
    #     #     definition = PyDictionary.PyDictionary.meaning(word)
    #     #     # if definition.get('Adverb'):
    #     #     #     continue
    #     #     #     answer = ask_to_add_word_to_list(word, definition.get('Adverb'), 'adverb')
    #     #     #     if answer:
    #     #     #         adverb.append(word.lower())
    #     #     if definition.get('Verb'):
    #     #         answer = ask_to_add_word_to_list(word, definition.get('Verb'), 'verb')
    #     #         if answer == "Skip":
    #     #             skippable.append(word.lower())
    #     #         elif answer:
    #     #             possitive_verb.append(word.lower())
    #     #             local_words.append(word)
    #     #         else:
    #     #             negative_verb.append(word.lower())
    #     #     if definition.get('Noun'):
    #     #         answer = ask_to_add_word_to_list(word, definition.get('Noun'), 'noun')
    #     #         if answer == "Skip":
    #     #             skippable.append(word.lower())
    #     #         elif answer:
    #     #             possitive_noun.append(word.lower())
    #     #             local_words.append(word)
    #     #         else:
    #     #             negative_noun.append(word.lower())
    #     #             local_words.append(word)
    #     #     if definition.get('Adjective'):
    #     #         answer = ask_to_add_word_to_list(word, definition.get('Adjective'), 'adjective')
    #     #         if answer == "Skip":
    #     #             skippable.append(word.lower())
    #     #         elif answer:
    #     #             possitive_adjective.append(word.lower())
    #     #             local_words.append(word)
    #     #         else:
    #     #             negative_adjective.append(word.lower())
    #     #     if definition.get('Adverb'):
    #     #         answer = ask_to_add_word_to_list(word, definition.get('Adverb'), 'adverb')
    #     #         if answer == "Skip":
    #     #             skippable.append(word.lower())
    #     #         elif answer:
    #     #             possitive_adverb.append(word.lower())
    #     #             local_words.append(word)
    #     #         else:
    #     #             negative_adverb.append(word.lower())
    #     else:
    #         print(word)
    #         retort = input('Add this to list?')
    #         # list_of_choices = ['negative_verb', 'negative_noun', 'negative_adverb', 'negative_adjective', 'possitive_verb', 'possitive_noun', 'possitive_adverb', 'possitive_adjective']
    #         if retort.lower() in ("n", "N", "No", "NO", "no"):
    #             continue
    #         elif retort.lower() in ("y", "Y", "Yes", "YES", "yes"):
    #             possitive_noun.append(retort.lower())
    #             local_words.append(retort.lower())
    #         elif retort.lower in ('skip', 'Skip'):
    #             skippable.append(retort.lower())
    #
            # if "bad" in retort.lower():
            #     list_of_choices = [negative_verb, negative_noun, negative_adverb, negative_adjective]
            # elif "good" in retort.lower():
            #     list_of_choices = [possitive_verb, possitive_noun, possitive_adverb, possitive_adjective]
            # # for key, value in list_of_choices.items():
            # #     print(key)
            # print('Which list would you like to add this to?')
            # for choice in list_of_choices:
            #     print(choice)
            # response = input()
            # if "" == response.lower():
            #     skippable.append(word)
            # if "skip" in response.lower():
            #     skippable.append(word)
            # if "conjunction" in response.lower():
            #     conjunctions.append(word)
            #     local_words.append(word)
            # if 'adverb' in response:
            #     if retort in ('b', 'B', ''):
            #         negative_adverb.append(word.lower())
            #     if retort in ('g', 'G'):
            #         possitive_adverb.append(word.lower())
            #         local_words.append(word)
            # elif 'verb' in response:
            #     if retort in ('b', 'B', ''):
            #         negative_verb.append(word.lower())
            #     if retort in ('g', 'G'):
            #         possitive_verb.append(word.lower())
            #         local_words.append(word)
            # if 'noun' in response:
            #     if retort in ('b', 'B', ''):
            #         negative_noun.append(word.lower())
            #     if retort in ('g', 'G'):
            #         possitive_noun.append(word.lower())
            # if 'adjective' in response:
            #     if retort in ('b', 'B', ''):
            #         negative_adjective.append(word.lower())
            #     if retort in ('g', 'G'):
            #         possitive_adjective.append(word.lower())
            #         local_words.append(word)
        # massive_list.append(word.lower())
    # json_file = json.dumps(keywords)
    # file = open('keywords.txt', 'r+')
    # file.write(json_file)
    # file.close()
    create_resume(local_words, company, job_title, location, url)


def select_sentence():
    sentences = {"1": "Used Python, Javascript, and other Languages to develop and maintain API through AWS Gateway",
    "2": "Implement SCRUM to prioritize the development and maintenance of functions in AWS Lambda.",
    "3": "Communicate with internal and external teams to create and consume API endpoints.",
    "4": "Work closely with team members to design scalable databases that are free of monolithic architecture.",
    "5": "Update git repositories to ensure that work is not being duplicated or lost. ",
    "6": "Researched and implemented new services and standards to ensure the fastest and most legible code.",
    "7": "Custom"}
    for key, value in sentences.items():
        print(key, value)
    while True:
        choice = input('Which sentence would you like to use?')
        if choice == "6":
            sentences["6"] = input("Please type your sentence.")
        if sentences.get(choice):
            return sentences[choice]
        else:
            print('Invalid Input')




def create_resume(local_words, company, job_title, location, url):
    job_sentence = []
    print(local_words)
    if len(job_sentence) <= 4:
        while len(job_sentence) < 4:
            job_sentence.append(select_sentence())
    from docx.shared import Inches

    document = Document()
    Pt = docx.shared.Pt
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(10)
    font.name = 'Garamond'

    header_section = document.sections[0].header
    name = header_section.paragraphs[0].add_run('SHELBY WHITTAKER')
    name.font.size = Pt(21)
    header_section.alignment = 'center'
    info = header_section.add_paragraph().add_run('2060 EAST DORIS WAY • HOLLADAY, UT  84124')
    # info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.font.size = Pt(7.5)
    info_pt_2  = header_section.add_paragraph().add_run('PHONE 801-598-7752 • E-MAIL LOUWHI1992@HOTMAIL.COM')
    info_pt_2.font.size = Pt(7.5)
    # info_pt_2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph_format = document.styles['Normal'].paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    objective_header = document.add_paragraph().add_run('OBJECTIVE')
    objective_header.font.size = docx.shared.Pt(12)
    objective_header.font.underline = WD_UNDERLINE.SINGLE
    table = document.add_table(rows=0, cols=2)
    row_cells = table.add_row().cells
    row_cells[0].width = Inches(1)
    row_cells[1].width = Inches(5)
    if job_title[:1] in ['a', 'e', 'i', 'o', 'u']:
        conjunction_word = "an"
    else:
        conjunction_word = 'a'
    SKILL = input(f'Fill in the following:\nTo design and maintain software as {conjunction_word} {job_title} so that {company} ')
    if SKILL == '':
        SKILL = f'provide lasting service to its customers.'
    objective_text = f'To design and maintain software as {conjunction_word} {job_title} so that {company} can {SKILL} '
    objective_invalid = True
    while objective_invalid:
        print(objective_text)
        objective_verification = input("Would you like to edit this Objective?")
        if objective_verification == "y":
            term_query = input("Which would you like to update? \n1) Job Title\n2) Company \n3) Skill")
            if term_query == "1":
                print(job_title)
                job_title = input('New Job Title:')
                objective_text = f'To design and maintain software as {conjunction_word} {job_title.strip()} so that {company} can {SKILL}'
            elif term_query == "2":
                print(company)
                company = input('New Company Name:')
                objective_text = f'To design and maintain software as {conjunction_word} {job_title.strip()} so that {company} can {SKILL}'
            elif term_query == "3":
                print(SKILL)
                objective_text = f'To design and maintain software as {conjunction_word} {job_title.strip()} so that {company} can {SKILL}'
        elif objective_verification == "n":
            objective_invalid = False
    objective_text = f'To design and maintain software as {conjunction_word} {job_title.strip()} so that {company} can {SKILL}'
    objective = row_cells[1].paragraphs[0].add_run(objective_text)
    objective.font.size = docx.shared.Pt(11)
    education_header = document.add_paragraph().add_run('EDUCATION')
    education_header.font.underline = WD_UNDERLINE.SINGLE
    education_header.font.size = docx.shared.Pt(12)
    # font.bold = True
    records = (
        ('', ['American Public University', "Bachelors of Information Technology", "Charles Town, West Virginia"], 'Present'),
        ('', ['American Public University', 'Associates of Computer Applications', 'Charles Town, West Virginia'], 'May 2018'),
        ('', ['Olympus High School', 'High School Diploma', 'Salt Lake City, Utah'], 'Class of 2010')
    )
    table = document.add_table(rows=0, cols=3)

    for column1, column2, column3 in records:
        row_cells = table.add_row().cells
        cell1 = row_cells[0]
        cell2 = row_cells[1]
        cell3 = row_cells[2]
        cell1.width = Inches(1)
        cell1.add_paragraph().add_run(' ').font.size = docx.shared.Pt(4)
        run = cell2.paragraphs[0].add_run(column2[0])
        run.bold = True
        run.font.size = docx.shared.Pt(12)
        cell2.add_paragraph().add_run(column2[1]).font.size = docx.shared.Pt(12)
        cell2.add_paragraph().add_run(column2[2]).font.size = docx.shared.Pt(10)
        cell2.width = Inches(4)
        cell2.add_paragraph('').add_run(' ').font.size = docx.shared.Pt(4)
        cell3.paragraphs[0].add_run(column3).font.size = docx.shared.Pt(12)
        cell3.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        cell3.width = Inches(1.5)
    knowledge_header = document.add_paragraph().add_run('KNOWLEDGE BASE')
    knowledge_header.font.size = Pt(11)
    knowledge_header.font.underline = WD_UNDERLINE.SINGLE
    records = (
        ('', ['Languages', 'Python', 'SQL', 'Javascript', 'PHP', 'CSS', 'HTML']),
        ('', ['Other Skills', 'AWS', 'APIs', 'Cloud Computing', 'Networking', 'Communication', 'Authentication'])
    )
    table = document.add_table(rows=0, cols=4)
    for column1, column2 in records:
        row_cells = table.add_row().cells
        cell1 = row_cells[0]
        cell2 = row_cells[1]
        cell3 = row_cells[2]
        cell4 = row_cells[3]
        cell1.width = Inches(1)
        cell2.width = Inches(2)
        cell3.width = Inches(2)
        cell4.width = Inches(2)
        title = cell2.paragraphs[0].add_run(column2[0])
        title.bold = True
        title.font.size = Pt(12)
        cell2.add_paragraph('', style='List Bullet').add_run(column2[4]).font.size = Pt(12)
        cell2.add_paragraph('', style='List Bullet').add_run(column2[5]).font.size = Pt(12)
        cell3.add_paragraph('', style='List Bullet').add_run(column2[6]).font.size = Pt(12)
        cell3.add_paragraph('', style='List Bullet').add_run(column2[1]).font.size = Pt(12)
        cell4.add_paragraph('', style='List Bullet').add_run(column2[2]).font.size = Pt(12)
        cell4.add_paragraph('', style='List Bullet').add_run(column2[3]).font.size = Pt(12)
    work_history_header = document.add_paragraph().add_run('WORK HISTORY')
    work_history_header.font.size = Pt(11)
    work_history_header.font.underline = WD_UNDERLINE.SINGLE
    table = document.add_table(rows=6, cols=2)
    cellA = table.rows[0].cells[0]
    cellA.paragraphs[0].add_run('LiveView Technologies').bold = True
    cellA.paragraphs[0].style.font.size = Pt(12)
    cellB = table.rows[0].cells[1]
    cellC = table.rows[1].cells[0].merge(table.rows[1].cells[1])
    cellE = table.rows[2].cells[0]
    cellF = table.rows[2].cells[1]
    cellG = table.rows[3].cells[0].merge(table.rows[3].cells[1])
    cellI = table.rows[4].cells[0]
    cellJ = table.rows[4].cells[1]
    cellK = table.rows[5].cells[0].merge(table.rows[5].cells[1])

    cellB.paragraphs[0].add_run("April 2018 - Present").italics = True
    cellB.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    cellB.paragraphs[0].style.font.size = Pt(12)
    cellC.paragraphs[0].add_run(job_sentence[0])
    cellC.paragraphs[0].style = "List Bullet"
    cellC.add_paragraph(style="List Bullet").add_run(job_sentence[1])
    cellC.add_paragraph(style="List Bullet").add_run(job_sentence[2])
    cellC.add_paragraph(style="List Bullet").add_run(job_sentence[3])
    cellC.paragraphs[0].style.font.size = Pt(12)

    cellE.paragraphs[0].add_run("Plansource").bold = True
    cellE.paragraphs[0].style.font.size = Pt(12)
    cellF.paragraphs[0].add_run("2013-2016").italics = True
    cellF.paragraphs[0].style.font.size = Pt(12)
    cellF.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    cellI.paragraphs[0].add_run("Xerox").bold = True
    cellI.paragraphs[0].style.font.size = Pt(12)
    cellJ.paragraphs[0].add_run("July 2014 - December 2014").italics = True
    cellJ.paragraphs[0].style.font.size = Pt(12)
    cellJ.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    cellG.paragraphs[0].add_run('Obtained outlines of how data should be presented per Client request.')
    cellG.paragraphs[0].style = "List Bullet"
    cellG.add_paragraph(style="List Bullet").add_run('Consulted EDI analysts to resolve data issues.')
    cellG.add_paragraph(style="List Bullet").add_run('Provided suggestions on making systems more user friendly.')
    cellG.paragraphs[0].style.font.size = Pt(12)

    cellK.paragraphs[0].add_run('Enrolled Subscribers for Blue Shield of California')
    cellK.paragraphs[0].style = "List Bullet"
    cellK.add_paragraph(style="List Bullet").add_run('Processed applications for missing information')
    cellK.add_paragraph(style="List Bullet").add_run('Managed multiple systems based on customer needs')
    cellK.paragraphs[0].style.font.size = Pt(12)
    saving = True
    while saving:
        try:
            document.save(f"{job_title} at {company}.docx")
            saving = False
        except:
            input(f'{job_title} at {company}.docx already exists. Please rename the document to try again.')


#

def get_html_content():
    link_file = open('link_list.txt', 'r')
    list_of_links = link_file.read()
    list_of_links = json.loads(list_of_links)
    link_file.close()
    driver = webdriver.Chrome('chromedriver_win32/chromedriver.exe')  # Optional argument, if not specified will search path.
    driver.get('https://www.smithsfoodanddrug.com/search?query=toilet%20paper&searchType=natural&fulfillment=all');
    search_box = driver.find_element_by_id('searchbutton')
    time.sleep(5)
    search_box.send_keys('toiletpaper')
    search_box.submit()
    pages = True
    current_url = driver.current_url
    searching_url = current_url
    job_count = requests.get(searching_url).text
    soup = BeautifulSoup(job_count, 'html.parser')
    all_urls = soup.find_all('a')
    list_of_job_urls = []
    for url in all_urls:
        if url.get('href'):
            value = urzl.get('href')
            if re.match(r'/pagead.*', value) or re.match(r'/rc/.*', value):
                newline = f'https://www.indeed.com{value}'
                list_of_job_urls.append(newline)
    # file.write(json_file)
    # file.close()
    # create_resume(local_words)

    for link in list_of_job_urls:
        if link in list_of_links:
            continue
        else:
            job_description = requests.get(link).text
            soup = BeautifulSoup(job_description, 'html.parser')
            #Find Job Title
            header = str(soup.find('title'))
            header = header.replace('<title>', '')
            job_title = header.split('-')[0].strip()
            print(f"Job Title is {job_title}")
            #Find Company
            company_searching = soup.find_all('a')
            company = company_searching[9]
            if re.search(r'indeed\.com/cmp/.+\?', str(company)):
                company = re.findall('indeed\.com/cmp/.+\?', str(company))[0][15:-1]
            elif re.search(r'indeed\.com/viewjob\?cmp.+&?', link):
                company = re.findall(r'indeed\.com/viewjob\?cmp=.+&?', link)[0][15:-1]
            elif re.search(r'\..+company', str(company_searching)):
                company = re.findall(r'\..+company', str(company_searching))
            else:
                company = 'Unknown'
            print(f"Company is {company}")
            #Find Location
            LOCATION = header.split('-')[1]
            print(f"Location is {LOCATION}")
            #Find Job Description
            paragraph = soup.find_all('p')
            para_graph = ''
            for url in paragraph:
                para_graph += str(url)
            paragraph_fix = re.findall('<.[^<>]*>?', para_graph)
            for fix in paragraph_fix:
                para_graph = para_graph.replace(fix, ' ')
                para_graph = para_graph.replace('  ', ' ')
            # print(para_graph)
            list_of_sentences = para_graph.split('.')
            list_of_words = list(dict.fromkeys(para_graph.split(' ')))
            print(link)
            determine_resume_content(list_of_sentences, list_of_words, company, job_title, LOCATION, link)
            list_of_links[link] = link
            json_file_list = json.dumps(list_of_links)
            link_file = open('link_list.txt', 'r+')
            link_file.write(json_file_list)
            link_file.close()
            validate_resume(f"{job_title} at {company}.docx", link, driver)
            submit_resume(link, driver)# TODO Remove this line
    link_file.close()


if __name__ == "__main__":
   get_html_content()
